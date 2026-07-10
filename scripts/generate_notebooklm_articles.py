#!/usr/bin/env python3
"""Generate NotebookLM-ready Lumi Social Intelligence article DOCX files.

This script writes Markdown sources and real .docx files using python-docx,
then verifies the generated DOCX packages by reading back their XML.
"""

from __future__ import annotations

from pathlib import Path
import re
import zipfile

from docx import Document  # type: ignore[import-untyped]
from docx.shared import Pt  # type: ignore[import-untyped]

ROOT = Path(__file__).resolve().parents[1]
ARTICLES = ROOT / "docs" / "articles"
ARTICLES.mkdir(parents=True, exist_ok=True)

REFERENCES_MD = r"""# Technical References for Lumi Social Intelligence

This source pack supports the NotebookLM articles on **Lumi Social Intelligence**, especially the technical article on **Lumi Layered Memory**, **Nuances**, and **Presence**. The references are used as grounding lanes, not as claims that the project implements a specific paper exactly.

## Agent memory, RAG, and personalization

1. **Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks** — Patrick Lewis et al., 2020. arXiv:2005.11401. https://arxiv.org/abs/2005.11401
   Supports the distinction between retrieval as a capability and Lumi Social Intelligence’s additional need for reviewable memory use, contextual appraisal, and action governance.

2. **Reflexion: Language Agents with Verbal Reinforcement Learning** — Noah Shinn et al., 2023. arXiv:2303.11366. https://arxiv.org/abs/2303.11366
   Relevant to agent self-reflection and learning loops; Lumi’s boundary is that social learning should route through review and receipts before becoming durable behavior.

3. **Generative Agents: Interactive Simulacra of Human Behavior** — Joon Sung Park et al., 2023. arXiv:2304.03442. https://arxiv.org/abs/2304.03442
   Useful comparison point for memory, reflection, and planning in agent behavior, while Lumi focuses on personal-agent boundaries, consent, and restraint.

4. **MemGPT: Towards LLMs as Operating Systems** — Charles Packer et al., 2023. arXiv:2310.08560. https://arxiv.org/abs/2310.08560
   Supports the idea that long-running agents need explicit memory management rather than relying only on context windows.

5. **Toolformer: Language Models Can Teach Themselves to Use Tools** — Timo Schick et al., 2023. arXiv:2302.04761. https://arxiv.org/abs/2302.04761
   Useful background for tool-using agents. Lumi’s contribution is not tool use itself, but the social decision layer around when an agent should act, ask, wait, or repair.

6. **ReAct: Synergizing Reasoning and Acting in Language Models** — Shunyu Yao et al., 2022. arXiv:2210.03629. https://arxiv.org/abs/2210.03629
   Background for reasoning/action loops. Presence can be understood as a social and safety gate around such loops.

## Human-agent interaction, mixed initiative, trust, and explainability

7. **Guidelines for Human-AI Interaction** — Saleema Amershi et al., 2019. ACM CHI. DOI:10.1145/3290605.3300233. https://doi.org/10.1145/3290605.3300233
   Supports the need for interaction norms such as showing uncertainty, allowing correction, and supporting graceful recovery.

8. **Mixed-Initiative User Interfaces** — Eric Horvitz, 1999. CHI. DOI:10.1145/302979.303030. https://doi.org/10.1145/302979.303030
   Relevant to Lumi’s Presence layer: initiative should be governed, contextual, and interruptible, not treated as inherently good.

9. **Some Requirements for Human-like Robots: Why the Recent Over-Emphasis on Autonomy and Emotion is a Distraction** — Kerstin Dautenhahn, 2007. Philosophical Transactions of the Royal Society B. DOI:10.1098/rstb.2007.2059. https://doi.org/10.1098/rstb.2007.2059
   Useful caution against equating social intelligence with emotion-performance or autonomy theater.

10. **Why Should I Trust You? Explaining the Predictions of Any Classifier** — Marco Tulio Ribeiro, Sameer Singh, Carlos Guestrin, 2016. DOI:10.1145/2939672.2939778. https://doi.org/10.1145/2939672.2939778
   Background for explainability. Lumi applies the same spirit to memory and action: show receipts and reasons, not hidden drift.

## Consent, privacy, and human-centered AI

11. **Privacy as Contextual Integrity** — Helen Nissenbaum, 2004. Washington Law Review. https://crypto.stanford.edu/portia/papers/RevnissenbaumDTP31.pdf
   Supports the idea that whether information use is appropriate depends on context, relationship, purpose, and transmission norms — central to Nuances.

12. **Human-Centered AI** — Ben Shneiderman, 2020. Oxford University Press / related HCAI work. https://hcil.umd.edu/human-centered-ai/
   Supports a design stance where automation, human control, reliability, and accountability are balanced rather than collapsed into autonomy.

13. **The Consentful Tech Project** — Una Lee, Dann Toliver, and Allied Media Projects, 2017. https://www.consentfultech.io/
   Practical design lens for consent, agency, and respectful defaults; useful for memory review, correction, and revocation boundaries.

## Safety, evaluation, and release engineering

14. **Artificial Intelligence Risk Management Framework (AI RMF 1.0)** — NIST, 2023. https://www.nist.gov/itl/ai-risk-management-framework
   Supports the framing of risk management, measurement, governance, and documentation around AI systems.

15. **Red Teaming Language Models to Reduce Harms: Methods, Scaling Behaviors, and Lessons Learned** — Ganguli et al., 2022. arXiv:2209.07858. https://arxiv.org/abs/2209.07858
   Background for adversarial and anti-pattern testing; Lumi uses anti-pattern fixtures and release gates to catch unsafe social behavior before runtime.

16. **SLSA: Supply-chain Levels for Software Artifacts** — OpenSSF, specification v1.0. https://slsa.dev/spec/v1.0/
   Supports artifact provenance, build hygiene, and release discipline.

17. **Hermes Agent Documentation** — Nous Research, current docs. https://hermes-agent.nousresearch.com/docs
   Practical reference for the first planned host target, **Lumi for Hermes**. The articles should not expose private Hermes runtime state; they should discuss Hermes as a host runtime and adapter target.
"""

INTRO = r"""# Lumi Social Intelligence: Memory, Nuance, and Presence for AI Agents

Most AI assistants are becoming more capable in the obvious ways. They can use tools, search documents, write code, summarize meetings, schedule tasks, and carry more context than earlier systems could. Yet the moments that break trust are often not the moments where the assistant lacks raw capability. They are the moments where it remembers the wrong thing, uses the right memory at the wrong time, interrupts a fragile moment, turns a correction into a personality rule, or treats a human signal as permission to act.

That is the gap Lumi Social Intelligence is built for.

Lumi Social Intelligence is a **social-intelligence layer for agents**: adaptive memory, nuance, and presence with review, consent, and repair. Its purpose is simple to say and difficult to build: agents that remember carefully, read the room better, and know when not to act.

The project begins from a practical observation. Long-term agents do not only need bigger context windows. They need judgment around context. If an assistant stores everything but cannot explain why a memory matters, it becomes opaque. If it interprets every mood, correction, or preference as a permanent trait, it becomes invasive. If it is encouraged to be proactive without a strong restraint layer, it becomes noisy or socially clumsy. Better memory alone does not make an agent socially intelligent. It only gives the agent more material to misuse.

Lumi Social Intelligence separates the problem into three cooperating products:

```text
Lumi Layered Memory -> Nuances -> Presence
```

In plain language:

```text
Lumi Layered Memory gives context.
Nuances reads the moment.
Presence decides whether to speak, wait, ask, act, repair, or stay quiet.
```

This separation matters. Many agent systems blur memory, interpretation, and action into one invisible personalization loop. A user says something once, the assistant silently absorbs it, and future behavior changes without a clear receipt. That may feel magical in a demo, but over time it becomes difficult to audit. Did the assistant learn a durable preference? Did it misunderstand a temporary mood? Is it adapting to the person, or to a single stressed message? Should it act on the memory now, ask first, or leave it alone?

Lumi Social Intelligence treats those as different questions.

**Lumi Layered Memory** handles continuity. It is responsible for useful remembered context: preferences, durable facts, project state, corrections, and other information that should help the agent work with the user over time. But its public contract is not “remember everything.” It is “remember carefully.” Memory should be reviewable. Memory use should have receipts. Corrections and revocations should matter. A long-running assistant should not silently rewrite a person from stray signals.

**Nuances** handles moment-reading. A memory does not mean the same thing in every situation. A correction may be a stable preference, a one-time frustration, a tone boundary, or a signal that the assistant should slow down. A user’s short reply may mean they are busy, annoyed, focused, tired, joking, or simply concise. Nuances does not pretend to know all of that with certainty. Its job is to preserve ambiguity, notice relevant signals, and propose humble interpretations. It reads the moment around the memory before anything becomes durable behavior.

**Presence** handles governed initiative. This is where the system decides whether to speak, wait, ask, act, repair, or hold silence. Presence makes restraint a supported outcome. A socially intelligent agent is not one that always fills the room. Sometimes the best action is to wait. Sometimes it should ask before using a sensitive memory. Sometimes it should repair a mistake quickly and plainly. Sometimes it should stay out of the way.

The architecture is designed around review, consent, and repair. Those are not decorative ethics words placed on top of the product. They are part of the product shape. A memory layer without review becomes hidden profiling. A nuance layer without humility becomes mind-reading theater. A presence layer without restraint becomes a notification cannon in a nice coat.

The first release doorway is **Lumi Social Intelligence**. Development can happen in private workshops, but the public release surface should be clean, tested, documented, and safe to inspect. That release doorway exists to keep the public story coherent: three cooperating products working as one social-intelligence layer, rather than a scattered pile of internal experiments.

The first planned host target is **Lumi for Hermes**. Hermes Agent is a natural first home because it is built around a personal assistant model: tools, memory, scheduled work, desktop automation, and long-running collaboration. But Lumi Social Intelligence is designed as a host-neutral layer. Host adapters should be thin runtime bindings. The product idea remains the same: careful continuity, humble appraisal, governed initiative, visible review, and fail-closed safety.

The human origin of the project is part of why it looks different from a conventional agent framework. Lumi Social Intelligence was conceived by one person building a full social-intelligence layer for Hermes Agent: a non-developer with a project-manager mindset, QA discipline, and music-producer sensitivity to timing, tone, dynamics, and emotional arc. That origin matters because the failure modes were noticed less as abstract software bugs and more as lived product failures.

A project manager notices when a workflow has no owner, no gate, or no clear handoff. A QA thinker notices edge cases, regressions, confusing states, and the little horrors that only appear after repetition. A music producer notices timing, tension, release, dynamics, and when one extra element ruins the feeling. Combined, those instincts point to a different kind of agent problem: not “can the model produce a good response once?” but “can the assistant behave well over time without becoming invasive, flat, clingy, noisy, or opaque?”

That question changes the design. The goal is not to make an assistant sound warmer by default. Warmth without boundaries can become manipulation. The goal is not to make it more proactive by default. Proactivity without timing becomes interruption. The goal is not to make it remember more by default. Memory without context becomes a liability. Lumi Social Intelligence is about adaptive restraint as much as adaptive capability.

This is why the product promise is deliberately modest and serious:

> Agents that remember carefully, read the room better, and know when not to act.

It is not claiming universal emotional understanding. It is not claiming that an AI can infer a person’s inner life. It is not a persona pack. It is an architecture for making agent behavior more inspectable, more correctable, and more socially aware at the boundaries where memory, interpretation, and action meet.

For users, the practical value is continuity without creepiness. An assistant should remember the things that make collaboration smoother, but it should also show what it thinks it knows and allow correction. It should adapt to tone and context, but without converting every moment into a permanent psychological profile. It should be able to help at the right time, but it should also be able to leave a moment alone.

For agent builders, the value is a cleaner product model. Instead of treating “memory” as one feature and “proactivity” as another, Lumi Social Intelligence gives a three-part loop:

1. What context is relevant and reviewable?
2. What does the current moment suggest, with uncertainty preserved?
3. What should the assistant do, if anything?

That loop is small enough to test and strong enough to shape a real product. It makes room for synthetic fixtures, anti-pattern tests, release gates, review cards, and fail-closed behavior. It turns social intelligence from a vibe into something that can be inspected.

The early release path is intentionally cautious. The repository is a private playground until release gates, public documentation, installer artifacts, and privacy scans are ready. Private development repositories can remain messy and experimental. The public doorway should only receive curated, tested, public-safe material. Raw runs, private memories, diaries, local runtime state, scheduler internals, credentials, chat IDs, private coordinates, and unverified host claims do not belong in the release surface.

That boundary is not bureaucracy. It is the same philosophy applied to the project itself: remember carefully, interpret carefully, act carefully. The product cannot credibly ask agents to respect context if its own release process leaks private context. The architecture has to live its own values.

Lumi Social Intelligence is therefore not just an assistant feature. It is a way of thinking about long-running AI systems as social software. Once an agent persists across days, projects, moods, corrections, and personal workflows, it is no longer enough for it to be technically capable. It must be able to handle continuity without ownership, interpretation without arrogance, and initiative without entitlement.

That is the missing layer Lumi Social Intelligence is trying to build.
"""

TECH = r"""# Inside Lumi Social Intelligence: Layered Memory, Nuances, and Presence

Lumi Social Intelligence is a social-intelligence layer for agents: adaptive memory, nuance, and presence with review, consent, and repair. It is built from three cooperating products:

```text
Lumi Layered Memory -> Nuances -> Presence
```

The simplest technical description is:

```text
Lumi Layered Memory gives context.
Nuances reads the moment.
Presence decides whether to speak, wait, ask, act, repair, or stay quiet.
```

This architecture starts from a limitation in many agent systems. Retrieval, memory, and tool use can increase capability, but they do not automatically create social judgment. A retrieval-augmented system may find the right fact and still use it at the wrong time. A memory system may store useful preferences but fail to distinguish durable preference from temporary mood. A tool-using agent may be able to act, but still lack a principled reason to decide whether acting is appropriate.

Lumi Social Intelligence treats memory, interpretation, and action as separate layers that must cooperate under review. This separation is the main technical idea.

## 1. Why agents need more than retrieval

Retrieval-augmented generation showed that language models can be grounded with external knowledge instead of relying only on model weights. That is useful, but personal agents need a stricter problem statement. They are not merely retrieving public facts. They are handling continuity about a person, a workflow, projects, preferences, corrections, private constraints, and sometimes emotionally weighted context.

That changes the requirements. Personal memory needs provenance. It needs revocation. It needs a difference between “the user said this once” and “this is a stable preference.” It needs a way to show why a memory was used. It needs boundaries around what should never become public release material. It also needs humility, because a remembered sentence is not the same thing as a full understanding of a person.

Agent memory projects such as Generative Agents, MemGPT, Reflexion, ReAct, and related long-context or tool-using agent work provide useful background. They show how memory, reflection, planning, and action can improve agent behavior. Lumi Social Intelligence borrows from that larger research landscape but focuses on a narrower public product question: how should an agent remember, interpret, and act around a real user without creating hidden personalization, unwanted intimacy, noisy initiative, or unreviewable drift?

## 2. Layer one: Lumi Layered Memory

**Lumi Layered Memory** is the continuity layer. Its role is not to remember everything. Its role is to preserve useful continuity with clean boundaries.

The public contract is:

- durable memory is reviewable;
- memory use can be explained with receipts;
- corrections and revocations matter;
- private runtime state does not become public release material;
- identity, preferences, and durable user facts are not silently rewritten.

A layered memory system should distinguish between different kinds of remembered information. A stable user preference is not the same as a project note. A project note is not the same as a private emotional signal. A correction is not automatically a personality rule. A transient context cue should not become durable memory without a reason.

This is why Lumi Layered Memory is designed as the first layer rather than the whole system. Memory supplies context, but it does not decide alone what the current moment means. It should provide candidate continuity, citations, and receipts. It should make memory visible enough to review. But it should not silently mutate the assistant’s behavior whenever it sees a signal.

## 3. Layer two: Nuances

**Nuances** is the contextual appraisal layer. It sits between memory and action.

Nuances asks what a memory, correction, tone shift, uncertainty signal, or contextual cue might mean right now. Its job is not to mind-read. Its job is to preserve ambiguity while noticing useful signals. That distinction is important. Social intelligence in agents should not become fake certainty about a user’s inner state.

Nuances handles questions such as:

- Is this correction likely a durable preference or a one-time adjustment?
- Is the user asking for action, reflection, silence, or a small answer?
- Is a remembered fact relevant here, or would using it feel invasive?
- Is the assistant about to over-explain, over-comfort, or over-personalize?
- Is there enough evidence to adapt, or should the system ask first?
- Does the moment carry emotional weight that should slow the assistant down?

Nuances does not directly mutate durable memory. That is a key boundary. It may propose an interpretation, mark uncertainty, suggest a review card, or route a signal toward a later decision. But durable learning should remain inspectable. Otherwise, nuance becomes hidden profiling.

In research terms, Nuances connects to human-agent interaction, mixed-initiative systems, trust calibration, explainability, affective computing, and contextual privacy. But it intentionally avoids the trap of treating social intelligence as sentiment analysis. A sentiment score cannot tell an agent whether it should speak, ask, repair, or leave a moment alone. Nuances is about contextual appraisal under uncertainty, not emotion detection as a product promise.

## 4. Layer three: Presence

**Presence** is the governed initiative layer. It decides whether the assistant should speak, wait, ask, act, repair, or stay quiet.

Presence exists because capability creates temptation. Once an agent can remember, infer, schedule, send, automate, and run tools, the question is no longer only “can it?” The question becomes “should it, now, with this confidence, under these permissions, and with this evidence?”

Presence is the action and restraint gate. It should support outcomes like:

- speak with a short answer;
- ask a clarifying question;
- wait because the user is focused;
- suggest but do not act;
- prepare a review card;
- repair a prior mistake;
- stay quiet because the moment is not asking for intervention;
- fail closed because memory, configuration, confidence, or permission is insufficient.

Presence treats restraint as a feature. This is a product decision and a safety decision. Proactive agents often fail not because they lack capability, but because they lack a good model of interruption, consent, timing, and social cost. Presence makes those costs explicit.

## 5. Review-first architecture

The core governance principle is review first. Durable memory changes and proactive behavior should be inspectable before they become live behavior.

This can be implemented through review cards, synthetic fixtures, dry-run mode, release gates, and visible receipts. A review card can show what the system thinks it noticed, what evidence supports it, what uncertainty remains, and what action Presence would choose. That makes the system correctable. It also separates evaluation from live side effects.

A review-first architecture is especially important for social-intelligence work because the system deals with high-context signals. A bug in a calculator is visible. A bug in a social memory system can become a quiet behavioral drift. The user may not immediately know what changed or why. Review surfaces make the hidden layer visible.

## 6. Evaluation model

Lumi Social Intelligence should be tested with synthetic fixtures, anti-pattern tests, public/secret scans, and release gates.

Synthetic fixtures allow the project to test social behavior without exposing private chat logs or personal memories. Anti-pattern tests can cover overfamiliarity, unsupported inference, fake warmth, noisy initiative, memory overreach, unreviewed adaptation, and failure to repair. Release gates can verify that public-facing files use the correct product names, include required licensing, and avoid private/local material.

For 0.1.0, the system should remain review-gated and fail-closed. It should not introduce uncontrolled live autonomous behavior. A safe preview is not the same thing as a fully live social agent. The first release target is inspectability and installable structure, not dramatic autonomy.

The release doorway currently expects package/distribution names:

```text
lumi-layered-memory
lumi-nuances
lumi-presence
```

The first planned host-specific distribution is **Lumi for Hermes**.

## 7. Release architecture

Development happens in private repositories:

- **Lumi Layered Memory**
- **Nuances**
- **Presence**
- **Autoresearch**

Release-ready updates are promoted into **Lumi Social Intelligence**, the public-facing release doorway.

The promotion path is:

```text
Private repo passes release gate
→ export curated files via script
→ copy into Lumi Social Intelligence
→ run doorway release check
→ commit as release-ready update
```

This separation lets the private workshop remain experimental while the public doorway stays clean. It also protects against accidental exposure of raw runs, private memories, local runtime state, scheduler internals, credentials, chat IDs, coordinates, and unverified host claims.

**Autoresearch** remains private harness and evidence infrastructure. It is not the public product surface. That distinction is important: evidence-gathering machinery may contain internal experiments, local assumptions, evaluation data, and private iteration trails that should not be exposed as user-facing product.

## 8. Research context

Lumi Social Intelligence sits at the intersection of several research and engineering lanes.

From retrieval-augmented generation and memory-augmented agents, it takes the idea that models need external context and persistent state. From agent systems such as ReAct, Reflexion, MemGPT, and Generative Agents, it takes the observation that memory, reflection, and action loops can improve behavior over time. But it adds a product boundary: personal-agent memory should be reviewable, correctable, and context-aware.

From human-agent interaction and mixed-initiative research, it takes the importance of timing, initiative, uncertainty, and graceful handoff. Eric Horvitz’s mixed-initiative work is especially relevant because Presence is fundamentally a mixed-initiative gate. The assistant may initiate sometimes, but initiative must be governed.

From explainable AI and trust calibration, it takes the need for receipts. A user should be able to understand why a memory was used or why an action was proposed. Trust should not be produced by soothing language. It should be produced by visible behavior, correction paths, and recoverability.

From contextual integrity and consentful design, it takes the view that information use depends on context, purpose, and relationship. A fact being available does not mean it is appropriate to use. A memory being true does not mean it should be surfaced now. Nuances exists partly to handle that distinction.

From AI safety, red-teaming, and release engineering, it takes the habit of testing failure modes before runtime. Fail-closed behavior, anti-pattern fixtures, artifact hygiene, and public/secret scans are not separate from social intelligence. They are how social intelligence avoids becoming a nice story wrapped around unsafe automation.

## 9. Why this architecture is deliberately conservative

The conservative parts of Lumi Social Intelligence are not signs of weakness. They are the product.

A system that can remember a user over time should be slower to infer than a stateless chatbot. A system that can act should be more careful about permission than a system that only talks. A system that can adapt socially should show its reasoning and leave room for correction. The more personal the agent becomes, the more important it is that learning remains inspectable.

This is why the 0.1.0 scope is an installable Hermes preview with review gates, synthetic fixtures, and fail-closed behavior. The first milestone should prove the shape: memory provides context, Nuances appraises the moment, Presence governs initiative, and the release doorway keeps public artifacts clean.

Lumi Social Intelligence is not trying to make agents perform emotion. It is trying to make them safer and more useful at the boundary between memory and action. That is where many long-running assistants will either earn trust or lose it.

## References

The companion file `references.md` contains the technical reference pack used for this article, including work on retrieval-augmented generation, agent memory, mixed-initiative interaction, human-AI guidelines, contextual integrity, consentful design, AI risk management, red-teaming, supply-chain artifact hygiene, and Hermes Agent documentation.
"""

STORY = r"""# How Lumi Social Intelligence Was Created: From One-Person Concept to Social-Intelligence Layer for Hermes

Lumi Social Intelligence began with a simple but stubborn realization: the most interesting failures in a capable personal agent were no longer only technical failures.

A tool call could work. A memory could exist. A summary could be accurate. A scheduled reminder could fire. The agent could write, search, automate, explain, and keep track of projects. Yet something could still feel wrong. It might remember a detail but use it in the wrong emotional weather. It might be helpful but too loud. It might be warm but performative. It might turn a correction into a permanent rule. It might miss the difference between “do this now,” “think with me,” “leave me in flow,” and “just be short.”

Those were not solved by adding one more feature. They were social failures.

Lumi Social Intelligence was conceived by one man concepting a full social-intelligence layer for Hermes Agent: a non-developer, project manager, QA thinker, and music producer. That combination shaped the project from the beginning. It was not born as a clean software architecture diagram. It came from repeated observation of awkward agent behavior, pressure-testing of product boundaries, sensitivity to timing and tone, and a refusal to accept fake warmth as intelligence.

The non-developer part matters. A developer might naturally begin with code modules, APIs, and implementation constraints. This project began with behavior. What did the assistant do that felt useful? What did it do that felt off? When did memory help? When did memory become too much? When did proactivity feel caring, and when did it become interruption? When did a short reply preserve the moment better than a clever one?

The project-manager part matters too. A project manager sees handoffs, scope boundaries, ownership, gates, and the difference between a good idea and a releasable product. Lumi Social Intelligence did not stay at the level of “the assistant should be more emotionally intelligent.” It became a system of responsibilities: one layer for memory, one layer for interpretation, one layer for initiative. Each layer had to know what it owned and what it should not do.

The QA part may matter most of all. QA thinking does not trust the happy path. It looks for regressions, edge cases, bad assumptions, hidden states, confusing labels, and the little QA horrors that appear after repetition. A socially intelligent agent can fail in subtle ways: overfamiliarity, under-response, unsupported inference, memory drift, apology loops, over-explaining, fake empathy, intrusive proactivity, or cheerful action without permission. Those are test cases, not vibes.

And the music-producer part shaped the feeling of the architecture. Music production trains attention to timing, dynamics, tension, release, silence, texture, and emotional arc. One extra element can ruin a mix. One sound too early can break the groove. One overcompressed chorus can remove the life. The same is true in agent interaction. Sometimes the best contribution is not a bigger response but better timing. Sometimes presence means restraint. Sometimes silence is not absence; it is the right move.

From those instincts, the core separation emerged:

```text
Lumi Layered Memory -> Nuances -> Presence
```

At first, the problem could have been called “better memory.” But memory alone was not enough. The assistant might remember a preference and still use it poorly. It might know that a user likes concise replies but fail to notice when a deeper explanation is invited. It might know a project context but miss that the user is in a creative flow and does not want a management report. It might know a sensitive detail but not understand that using it in public-facing material would be wrong.

So memory became only the first layer.

**Lumi Layered Memory** was the continuity layer: useful facts, preferences, corrections, project context, and receipts. But it had to be reviewable. It had to avoid silent rewrites. It had to make durable memory different from temporary signal. It had to support correction and revocation. The goal was not “remember more.” The goal was “remember carefully.”

The next gap was interpretation. A memory is not self-explanatory. Human communication is full of signals that are local, partial, ambiguous, and timing-dependent. A correction might mean “change this forever,” or it might mean “not right now.” A joking phrase might be part of a shared tone, not a durable preference. A short message might be frustration, focus, tiredness, or simply efficiency. Treating every cue as certainty would be dangerous.

That became **Nuances**.

Nuances was designed to read the moment without pretending to own the truth of it. It appraises tone, correction, consent, uncertainty, emotional weight, and context shifts. It keeps inference humble. It can propose that something may matter, but it should not directly mutate durable memory on its own. That boundary came from seeing how easily personalization becomes hidden profiling when interpretation and storage are fused.

The final gap was action.

If an assistant has memory and can interpret signals, should it speak? Ask? Wait? Act? Prepare something? Repair? Stay quiet? This is the layer many systems skip. They treat proactivity as a feature by itself, as if more initiative automatically means a better assistant. In lived use, that is false. A proactive assistant without social governance becomes a tiny notification cannon in a blazer. Useful sometimes, yes. But often: too much.

That became **Presence**.

Presence is governed initiative. It decides whether the assistant should speak, wait, ask, act, repair, or stay quiet. It makes restraint a valid output. It treats timing as part of intelligence. It recognizes that the right action can be wrong if the moment is wrong. It gives the system a way to fail closed when confidence, permission, or context is insufficient.

The architecture was not invented in one dramatic moment. It formed through iteration: noticing failures, naming them, clustering them, testing proposed boundaries, rejecting vague “be more human” language, and turning lived observations into product surfaces that could be reviewed.

Research entered as a pressure test, not as decoration. The idea was compared against agent memory, retrieval-augmented generation, long-running agent systems, human-agent interaction, mixed-initiative interfaces, trust calibration, explainability, contextual integrity, consentful technology, affective computing, red-teaming, and release engineering. The project did not need to pretend that every piece was unprecedented. The stronger claim was more precise: these known lanes needed to be combined into a product layer focused on memory, moment-reading, and governed initiative for real personal agents.

Implementation followed the same separation. Private development repositories could contain the working mess: experiments, tests, drafts, fixtures, and evidence. But the public release needed a clean doorway. That doorway became **Lumi Social Intelligence**.

Development happens privately in:

- **Lumi Layered Memory**
- **Nuances**
- **Presence**
- **Autoresearch**

Tested, reviewed, public-safe updates are promoted into the **Lumi Social Intelligence** release repository. Autoresearch stays private as harness and evidence infrastructure, not as the public product. This matters because evidence work often contains raw runs, local context, internal tests, and private assumptions. The public surface should contain curated code, docs, examples, and release artifacts — not the workshop floor.

The promotion path is intentionally explicit:

```text
Private repo passes release gate
→ export curated files via script
→ copy into Lumi Social Intelligence
→ run doorway release check
→ commit as release-ready update
```

That release discipline is part of the story. Lumi Social Intelligence is not just about making agents polite. It is about making social adaptation testable. The project uses synthetic fixtures instead of private chat logs, anti-pattern tests instead of vibes, public/secret scans instead of hope, and fail-closed gates instead of uncontrolled live autonomy.

The early release target is **0.1.0: installable Hermes preview**. The distribution names are:

```text
lumi-layered-memory
lumi-nuances
lumi-presence
```

The first host-specific path is **Lumi for Hermes**. Hermes is the right first host because the project emerged from a real long-running Hermes assistant context: tools, memory, scheduled work, files, creative collaboration, and daily usefulness. But Lumi Social Intelligence is not merely a private personality tweak. The public product is the architecture: continuity, appraisal, initiative, review, consent, and repair.

Why was it created?

Because useful agents will increasingly live near human workflows. They will remember projects, preferences, constraints, mistakes, and patterns. They will have enough agency to schedule, send, write, automate, and interrupt. In that world, technical capability without social judgment becomes dangerous in small, ordinary ways. Not cinematic danger. Daily erosion danger. The assistant that remembers too much. The assistant that speaks at the wrong time. The assistant that assumes closeness. The assistant that silently adapts from a single moment. The assistant that optimizes productivity while damaging trust.

Lumi Social Intelligence exists to make those failures less likely.

It was created for people who need AI to work with them over time: creators, project managers, researchers, operators, high-context users, sensitive users, and anyone whose work depends on timing, tone, trust, and continuity. It is also for non-developers who can see product architecture before they can personally implement every line of code. One of the quiet arguments of the project is that deep product insight does not only come from formal engineering background. Sometimes it comes from watching a system fail in real life, naming the failures clearly, and refusing to let the machine hide them under charm.

The creation story should not be mistaken for lone-genius mythology. Lumi Social Intelligence was not magic. It was noticing, naming, sorting, testing, and building gates. It was PM structure, QA skepticism, music-person timing, and enough technical collaboration with Hermes to turn the idea into files, repos, release checks, and eventually installable artifacts.

The project’s most important principle may be this: warmth is not enough. Intelligence is not enough. Memory is not enough. An agent that works beside a person over time needs social boundaries it can actually follow.

That is what Lumi Social Intelligence is trying to provide.

A memory layer that can be inspected.

A nuance layer that stays humble.

A presence layer that knows when not to act.

And a release doorway clean enough to ship without betraying the private work that taught the system why those boundaries matter.
"""

FILES = {
    "01-introduction-lumi-social-intelligence": INTRO,
    "02-technical-lumi-layered-memory-nuances-presence": TECH,
    "03-creation-story-lumi-social-intelligence": STORY,
    "references": REFERENCES_MD,
}


def add_markdown_to_doc(doc: Document, text: str) -> None:
    in_code = False
    code_lines = []
    for raw in text.splitlines():
        line = raw.rstrip()
        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                in_code = False
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, line[2:])
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            add_runs(p, re.sub(r"^\d+\. ", "", line))
        elif line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            run = p.add_run(line[2:])
            run.italic = True
        else:
            p = doc.add_paragraph()
            add_runs(p, line)


def add_runs(p, text: str) -> None:
    # Minimal bold/inline-code support for Markdown-ish source.
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Courier New"
        else:
            p.add_run(part)


def write_docx(stem: str, text: str) -> Path:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(11)
    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)
    add_markdown_to_doc(doc, text)
    out = ARTICLES / f"{stem}.docx"
    doc.save(out)
    return out


def verify_docx(path: Path, required: list[str]) -> str:
    with zipfile.ZipFile(path) as z:
        bad = z.testzip()
        if bad:
            raise RuntimeError(f"bad zip member {bad} in {path}")
        xml = z.read("word/document.xml").decode("utf-8", errors="replace")
    # Verification only needs plain-text term presence; avoid parsing generated XML.
    full = re.sub(r"<[^>]+>", " ", xml)
    full = re.sub(r"\s+", " ", full)
    missing = [term for term in required if term not in full]
    if missing:
        raise RuntimeError(f"{path.name} missing terms: {missing}")
    return full[:300]


def main() -> None:
    outputs = []
    for stem, text in FILES.items():
        md = ARTICLES / f"{stem}.md"
        md.write_text(text, encoding="utf-8")
        docx = write_docx(stem, text)
        outputs.append(docx)

    checks = {
        "01-introduction-lumi-social-intelligence.docx": ["Lumi Social Intelligence", "Lumi Layered Memory", "Nuances", "Presence", "Hermes Agent", "social-intelligence layer for agents"],
        "02-technical-lumi-layered-memory-nuances-presence.docx": ["lumi-layered-memory", "lumi-nuances", "lumi-presence", "Autoresearch", "Private repo passes release gate"],
        "03-creation-story-lumi-social-intelligence.docx": ["one man concepting a full social-intelligence layer for Hermes Agent", "non-developer", "project manager", "QA", "music producer"],
        "references.docx": ["Retrieval-Augmented Generation", "Guidelines for Human-AI Interaction", "Privacy as Contextual Integrity", "Hermes Agent Documentation"],
    }
    for docx in outputs:
        verify_docx(docx, checks[docx.name])
    print("Generated and verified DOCX files:")
    for docx in outputs:
        print(docx)


if __name__ == "__main__":
    main()
